"""Bridged conversion — logic / boundary / error / object-state.

Ten formats have no native extractor and are read through a modern sibling.
Two thirds of what matters here can be asserted without the external program
present, and is: which formats route to a bridge, what a caller is told when
the program is absent, and that the result reports the original format rather
than the sibling's.

The remaining third needs LibreOffice actually installed, and skips when it is
not. Those tests are marked so a skip is visible rather than silently counted
as a pass.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from convilyn import local
from convilyn.local._engine.formats import DocumentFormat
from convilyn.local._engine.markdown.bridge import extract_via_bridge
from convilyn.local._run import ToolNotInstalledError
from convilyn.local._tools import LIBREOFFICE, find_tool
from convilyn.local.types import Engine

_HAS_LIBREOFFICE = find_tool(LIBREOFFICE) is not None

needs_libreoffice = pytest.mark.skipif(
    not _HAS_LIBREOFFICE,
    reason="LibreOffice is not installed on this machine",
)


@pytest.fixture
def odt_document(tmp_path: Path) -> Path:
    """A real .odt, produced by converting a .docx we build ourselves.

    Written rather than committed: a binary fixture in the tree would be one
    more thing shipped to the public mirror, and this costs one extra hop.
    """
    import docx

    from convilyn.local._run import run_libreoffice

    document = docx.Document()
    document.add_heading("Bridged Heading", level=1)
    document.add_paragraph("A paragraph that survives the hop.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Region"
    table.cell(0, 1).text = "Growth"
    table.cell(1, 0).text = "APAC"
    table.cell(1, 1).text = "12%"

    source = tmp_path / "source.docx"
    document.save(str(source))

    returncode, _out, err = run_libreoffice(str(source), str(tmp_path), "odt")
    assert returncode == 0, err
    produced = tmp_path / "source.odt"
    assert produced.is_file(), err
    return produced


# ── 1. Logic — a bridged format converts, with structure intact ──────


@needs_libreoffice
class TestBridgedConversion:
    def test_an_odt_converts(self, odt_document: Path) -> None:
        result = local.convert(odt_document, to="md")

        assert result.ok is True

    def test_headings_survive_the_hop(self, odt_document: Path) -> None:
        local.convert(odt_document, to="md")

        assert "# Bridged Heading" in odt_document.with_suffix(".md").read_text(encoding="utf-8")

    def test_tables_survive_the_hop(self, odt_document: Path) -> None:
        local.convert(odt_document, to="md")
        rendered = odt_document.with_suffix(".md").read_text(encoding="utf-8")

        assert "| Region | Growth |" in rendered

    def test_the_result_reports_the_original_format(self, odt_document: Path) -> None:
        """Not the sibling's.

        The sibling is an implementation detail of reading; a caller who saw
        `docx` for a file they handed over as `.odt` would have no way to tell
        the two apart.
        """
        result = local.convert(odt_document, to="md")

        assert result.source_format == "odt"

    def test_the_route_taken_is_disclosed(self, odt_document: Path) -> None:
        """Nothing about the hop is hidden — it is in the warnings."""
        result = local.convert(odt_document, to="md")

        assert any("bridged: read via docx" in w for w in result.warnings)

    def test_the_engine_is_reported_as_the_office_suite(self, odt_document: Path) -> None:
        result = local.convert(odt_document, to="md")

        assert result.engine == "office-suite"


# ── 2. Boundary — which formats route where, tool or no tool ─────────


class TestRouting:
    @pytest.mark.parametrize("source", ["doc", "odt", "rtf", "xls", "ods", "ppt", "odp"])
    def test_office_formats_route_to_the_office_suite(self, source: str) -> None:
        route = local.capabilities().can(source, "md")

        assert route.engine == "office-suite"

    @pytest.mark.parametrize("source", ["epub", "mobi", "azw3"])
    def test_ebook_formats_route_to_the_ebook_tool(self, source: str) -> None:
        route = local.capabilities().can(source, "md")

        assert route.engine == "ebook"

    @pytest.mark.parametrize("source", ["pdf", "docx", "pptx", "xlsx", "csv", "txt", "xml"])
    def test_native_formats_need_no_bridge(self, source: str) -> None:
        route = local.capabilities().can(source, "md")

        assert route.engine == "structured"


# ── 3. Error — what a caller is told, and how it is classified ───────


class TestErrors:
    def test_a_format_with_no_sibling_is_refused(self, tmp_path: Path) -> None:
        """Routed here in error — a caller bug, not a machine fact."""
        with pytest.raises(ValueError, match="no sibling format"):
            extract_via_bridge(tmp_path / "x.docx", DocumentFormat.DOCX)

    def test_a_missing_tool_is_a_missing_dependency_not_a_failed_conversion(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """The two are acted on differently, so they must not collapse.

        Reached only in a race — `capabilities()` saw the program and it was
        gone by the time the conversion started — but the classification is what
        tells a caller whether to install something or to look at the file.
        """
        source = tmp_path / "doc.odt"
        source.write_bytes(b"not really an odt")

        def gone(*_args: object, **_kwargs: object) -> None:
            raise ToolNotInstalledError("LibreOffice vanished")

        # Patched on `_runners`, which is where the bridge is now called from.
        # `api.convert` still translates the error, so the classification under
        # test is unchanged — only the module that invokes the bridge moved.
        monkeypatch.setattr(local._runners, "extract_via_bridge", gone)
        monkeypatch.setattr(
            local.api, "plan", lambda *a, **k: _available_route("odt", "office-suite")
        )

        with pytest.raises(local.MissingDependencyError):
            local.convert(source, to="md")


def _available_route(source_format: str, engine: Engine) -> local.Route:
    """A route that claims to be runnable, for the race above.

    ``engine`` is typed as ``Engine`` rather than ``str`` for the same reason
    ``_routes._requirements_for`` is: declared as ``str`` it needed a
    ``# type: ignore[arg-type]`` here, which suppressed the one check that catches
    a mistyped engine in a fixture.
    """
    return local.Route(
        source_format=source_format,
        target_format="md",
        engine=engine,
        available=True,
        requirements=(
            local.Requirement(
                kind="external_tool",
                name="libreoffice",
                available=True,
                install_hint="install it",
            ),
        ),
    )


# ── 4. Object-state — the projection is wired, not merely present ────


class TestTheBridgeIsReachable:
    def test_the_document_runner_uses_the_bridge_for_a_format_with_no_extractor(self) -> None:
        """Guard against the module existing but nothing calling it.

        The caller is `_runners.run_document`, not `api`: extraction moved there
        when engine dispatch became a table. The guard still answers the same
        question — is the generated bridge actually reached — and asserting it
        against `api` would now pass only by accident of an unused import.
        """
        from convilyn.local import _runners

        assert _runners.extract_via_bridge is extract_via_bridge

    def test_every_bridged_format_has_no_native_extractor(self) -> None:
        """Otherwise the bridge would never be the path taken."""
        from convilyn.local._engine.markdown.registry import bridge_for, extractor_for

        bridged = [f for f in DocumentFormat if bridge_for(f) is not None]

        assert bridged, "no format is bridged; every test above is vacuous"
        assert all(extractor_for(f) is None for f in bridged)


# ── 5. Isolation — the profile the external program runs against ─────


class TestLibreOfficeProfileIsolation:
    """LibreOffice serialises on its user profile, so we do not share theirs.

    An `soffice` started against a profile another instance holds exits
    non-zero having converted nothing, and says nothing useful about why. A
    user with LibreOffice open on their desktop hit that on their first
    conversion; this suite hit it intermittently against a leftover instance.
    """

    def test_the_invocation_names_a_profile_of_our_own(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        from convilyn.local import _run as run_module

        seen: list[list[str]] = []
        monkeypatch.setattr(run_module, "find_tool", lambda _tool: Path("soffice"))
        monkeypatch.setattr(
            run_module, "_run", lambda argv, timeout: seen.append(argv) or (0, "", "")
        )

        run_module.run_libreoffice(str(tmp_path / "a.odt"), str(tmp_path), "docx")

        assert seen, "the runner never reached the subprocess seam"
        assert any(arg.startswith("-env:UserInstallation=file:") for arg in seen[0])

    def test_the_profile_is_reused_rather_than_recreated(self) -> None:
        """A fresh profile costs LibreOffice a first-run initialisation.

        Paying that per call would charge every converted file for a collision
        that one directory per machine already prevents.
        """
        from convilyn.local._run import _private_profile

        assert _private_profile() == _private_profile()

    def test_the_profile_is_not_the_users_own(self) -> None:
        from convilyn.local._run import _private_profile

        assert "convilyn" in _private_profile()
