"""DOCX extraction, against real .docx files built with python-docx."""

from __future__ import annotations

import io
from pathlib import Path

import pytest

from convilyn.local._engine.markdown.docx import extract
from convilyn.local._engine.markdown.render import render


def _png(size: tuple[int, int] = (256, 256)) -> bytes:
    """A content-bearing PNG — a flat one would be filtered as decorative."""
    from PIL import Image

    image = Image.new("RGB", size)
    image.putdata(
        [
            ((x * 7) % 256, (y * 11) % 256, (x * y) % 256)
            for y in range(size[1])
            for x in range(size[0])
        ]
    )
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


@pytest.fixture
def styled_docx(tmp_path: Path) -> Path:
    import docx

    document = docx.Document()
    document.add_heading("Chapter One", level=1)
    document.add_paragraph("An ordinary paragraph of prose.")
    document.add_paragraph("first bullet", style="List Bullet")
    document.add_paragraph("second bullet", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "name"
    table.cell(0, 1).text = "qty"
    table.cell(1, 0).text = "bolt"
    table.cell(1, 1).text = "4"
    document.add_heading("Sub-section", level=2)

    path = tmp_path / "s.docx"
    document.save(str(path))
    return path


class TestStyles:
    def test_heading_style_becomes_a_heading_block(self, styled_docx):
        headings = [b for b in extract(styled_docx).blocks if b.kind == "heading"]

        assert "Chapter One" in [b.text for b in headings]

    def test_heading_level_is_carried_through(self, styled_docx):
        levels = {b.text: b.level for b in extract(styled_docx).blocks if b.kind == "heading"}

        assert levels["Sub-section"] == 2

    def test_bullet_style_becomes_a_list_item(self, styled_docx):
        items = [b for b in extract(styled_docx).blocks if b.kind == "list_item"]

        assert len(items) == 2

    def test_plain_paragraph_stays_a_paragraph(self, styled_docx):
        texts = [b.text for b in extract(styled_docx).blocks if b.kind == "paragraph"]

        assert "An ordinary paragraph of prose." in texts

    def test_table_is_extracted_with_its_header(self, styled_docx):
        tables = [b for b in extract(styled_docx).blocks if b.kind == "table"]

        assert tables[0].rows[0] == ("name", "qty")


class TestDocumentOrder:
    def test_blocks_come_back_in_authored_order(self, styled_docx):
        """``document.paragraphs`` + ``document.tables`` are two separate lists,
        so reading them would put the table AFTER the trailing heading. Only
        walking the body's XML children preserves the interleaving."""
        kinds = [b.kind for b in extract(styled_docx).blocks]

        assert kinds == [
            "heading",
            "paragraph",
            "list_item",
            "list_item",
            "table",
            "heading",
        ]


class TestImages:
    @pytest.fixture
    def docx_with_image(self, tmp_path: Path) -> Path:
        import docx

        document = docx.Document()
        document.add_paragraph("Before the figure.")
        image_path = tmp_path / "fig.png"
        image_path.write_bytes(_png())
        document.add_picture(str(image_path))
        document.add_paragraph("After the figure.")

        path = tmp_path / "i.docx"
        document.save(str(path))
        return path

    def test_embedded_image_is_extracted(self, docx_with_image):
        images = [b for b in extract(docx_with_image).blocks if b.kind == "image"]

        assert len(images) == 1

    def test_image_carries_real_bytes(self, docx_with_image):
        """The old path wrote images to a temp dir nothing uploaded."""
        image = next(b.image for b in extract(docx_with_image).blocks if b.kind == "image")

        assert image is not None and image.data.startswith(b"\x89PNG")

    def test_image_media_type_comes_from_the_part(self, docx_with_image):
        image = next(b.image for b in extract(docx_with_image).blocks if b.kind == "image")

        assert image is not None and image.media_type == "image/png"

    def test_rendered_link_points_into_the_asset_dir(self, docx_with_image):
        assert "](assets/img-0001.png)" in render(extract(docx_with_image))

    def test_repeated_image_collapses_to_one_asset(self, tmp_path):
        """The same picture twice is one file and one description, not two."""
        import docx

        document = docx.Document()
        image_path = tmp_path / "fig.png"
        image_path.write_bytes(_png())
        document.add_picture(str(image_path))
        document.add_picture(str(image_path))
        path = tmp_path / "r.docx"
        document.save(str(path))

        assert len(extract(path).unique_images) == 1


class TestRobustness:
    def test_empty_document_produces_no_blocks(self, tmp_path):
        import docx

        path = tmp_path / "e.docx"
        docx.Document().save(str(path))

        assert extract(path).blocks == ()

    def test_source_format_is_reported(self, styled_docx):
        assert extract(styled_docx).source_format == "docx"
